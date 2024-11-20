from docx import Document

# Create a new Word Document
doc = Document()
doc.add_heading('Dataset for DH.ARC - Digital Humanities Advanced Research Centre', level=1)

# Adding Nodes Table
doc.add_heading('1. Nodes', level=2)
nodes_table = doc.add_table(rows=1, cols=6)
nodes_table.style = 'Table Grid'

# Adding Header
hdr_cells = nodes_table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Research Field'
hdr_cells[3].text = 'Position'
hdr_cells[4].text = 'Affiliation'
hdr_cells[5].text = 'Publication Count'

# Adding data to Nodes table
nodes_data = [
    ['R001', 'Silvio Peroni', 'Semantic Web', 'Professor', 'DH.ARC Researcher', '45'],
    ['R002', 'Francesca Tomasi', 'Digital Scholarly Editing', 'Professor', 'DH.ARC Researcher', '30'],
    ['R003', 'Fabio Vitali', 'Web Technologies', 'Professor', 'DH.ARC Researcher', '40'],
    ['R004', 'Marilena Daquino', 'Knowledge Representation', 'Postdoc', 'DH.ARC Researcher', '20'],
    ['R005', 'Paola Italia', 'Textual Scholarship', 'Professor', 'DH.ARC Researcher', '35'],
    ['R006', 'Sofia Pescarin', 'Virtual Reality in CH', 'Researcher', 'DH.ARC Researcher', '25'],
    ['R007', 'Giovanni Colavizza', 'Computational Humanities', 'Lecturer', 'DH.ARC Researcher', '15'],
]

for row_data in nodes_data:
    row_cells = nodes_table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Adding Edges Table
doc.add_heading('2. Edges', level=2)
edges_table = doc.add_table(rows=1, cols=5)
edges_table.style = 'Table Grid'

# Adding Header
hdr_cells = edges_table.rows[0].cells
hdr_cells[0].text = 'Source'
hdr_cells[1].text = 'Target'
hdr_cells[2].text = 'Collaboration Type'
hdr_cells[3].text = 'Project Name'
hdr_cells[4].text = 'Year'

# Adding data to Edges table
edges_data = [
    ['R001', 'R002', 'Co-authorship', 'LOD for Humanities', '2020'],
    ['R001', 'R003', 'Joint Project', 'Knowledge Graphs', '2021'],
    ['R002', 'R004', 'Supervision', 'Editing in Digital Age', '2022'],
    ['R003', 'R005', 'Co-authorship', 'Semantic Annotation', '2019'],
    ['R006', 'R007', 'Joint Project', 'VR in CH', '2023'],
    ['R004', 'R006', 'Co-authorship', 'Digital Knowledge', '2022'],
    ['R005', 'R007', 'Joint Project', 'Textual Analysis', '2020'],
]

for row_data in edges_data:
    row_cells = edges_table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Adding Optional Publications Table
doc.add_heading('3. Publications (Optional)', level=2)
publications_table = doc.add_table(rows=1, cols=5)
publications_table.style = 'Table Grid'

# Adding Header
hdr_cells = publications_table.rows[0].cells
hdr_cells[0].text = 'Publication ID'
hdr_cells[1].text = 'Title'
hdr_cells[2].text = 'Year'
hdr_cells[3].text = 'Authors'
hdr_cells[4].text = 'Field'

# Adding data to Publications table
publications_data = [
    ['P001', 'LOD for Humanities', '2020', 'R001, R002', 'Linked Open Data'],
    ['P002', 'Semantic Annotation', '2019', 'R003, R005', 'Semantic Web'],
    ['P003', 'Editing in Digital Age', '2022', 'R002, R004', 'Digital Scholarly Editing'],
]

for row_data in publications_data:
    row_cells = publications_table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Save the document
file_path = 'C:/Users/Lucrezia/OneDrive - Alma Mater Studiorum Università di Bologna/Network Analysis/Project/DHARC_Dataset.docx'
doc.save(file_path)

file_path
